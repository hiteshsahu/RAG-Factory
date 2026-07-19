# Copyright 2026 Hitesh Kumar Sahu — https://hiteshsahu.com
# SPDX-License-Identifier: Apache-2.0

from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path

from docx import Document
from raginator.core import Ingestor, RawDocument


class DocxIngestor(Ingestor):
    """Reads every .docx file under a directory, extracting paragraph text (The Suck-Inator)."""

    def __init__(self, source_dir: str | Path) -> None:
        self._source_dir = Path(source_dir)

    def ingest(self) -> Iterable[RawDocument]:
        for path in sorted(self._source_dir.glob("*.docx")):
            document = Document(str(path))
            paragraphs = [p.text for p in document.paragraphs if p.text]
            content = "\n\n".join(paragraphs)
            yield RawDocument(
                content=content,
                metadata={"path": str(path), "num_paragraphs": len(paragraphs)},
                source_id=str(path),
            )
